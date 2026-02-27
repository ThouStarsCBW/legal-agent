from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
import os
import time

app = Flask(__name__)
# 允许跨域，前端才能顺利调用接口
CORS(app)


# ==========================================
# 1. 法律问题问答接口 (POST)
# ==========================================
@app.route('/api/chat', methods=['POST'])
def chat():
    data = request.json
    user_message = data.get('message', '')

    # 这里预留给你们接入【腾讯元宝】或【腾讯元器】的 API
    time.sleep(1)  # 模拟大模型思考延迟
    mock_reply = f"关于“{user_message}”，根据《民法典》相关规定，建议您收集好微信聊天记录、转账凭证等证据。(注：此为后端模拟返回)"

    return jsonify({"code": 200, "data": {"reply": mock_reply}})


# ==========================================
# 2. 知识库检索接口 (GET)
# ==========================================
@app.route('/api/search', methods=['GET'])
def search():
    keyword = request.args.get('keyword', '')

    # [cite_start]根据赛题手册，这里预留接入得理提供的“开放平台API” [cite: 522]
    time.sleep(0.5)
    mock_results = [
        {"title": f"与“{keyword}”相关的法条：《中华人民共和国劳动合同法》第八十二条",
         "content": "用人单位自用工之日起超过一个月不满一年未与劳动者订立书面劳动合同的，应当向劳动者每月支付二倍的工资。"},
        {"title": f"相关案例参考：(2023)京01民终xxx号", "content": "法院认为，劳动者有权要求..."}
    ]

    return jsonify({"code": 200, "data": {"results": mock_results}})


# ==========================================
# 3. 法律文件识别接口 (POST)
# ==========================================
@app.route('/api/ocr', methods=['POST'])
def ocr():
    if 'file' not in request.files:
        return jsonify({"code": 400, "msg": "未上传文件"})

    file = request.files['file']
    # [cite_start]你们团队后续可以在这里调用 PaddleOCR 提取文字，再喂给小理AI分析 [cite: 521]
    time.sleep(1.5)

    return jsonify({
        "code": 200,
        "data": {
            "text": f"提取到的【{file.filename}】文本内容：\n甲方：张三\n乙方：李四\n违约金：200万元。",
            "risk": "风险提示：违约金比例可能过高，超出实际损失的30%法定保护上限。"
        }
    })


# ==========================================
# 4. 法律文书生成接口 (POST)
# ==========================================
@app.route('/api/generate', methods=['POST'])
def generate():
    data = request.json
    plaintiff = data.get('plaintiff', '原告')
    defendant = data.get('defendant', '被告')
    demand = data.get('demand', '诉讼请求')

    if not os.path.exists('temp_docs'):
        os.makedirs('temp_docs')

    file_path = f"temp_docs/{plaintiff}_起诉状.docx"
    doc = Document()
    doc.add_heading('民事起诉状', 0)
    doc.add_paragraph(f"原告：{plaintiff}")
    doc.add_paragraph(f"被告：{defendant}")
    doc.add_heading('诉讼请求：', level=1)
    doc.add_paragraph(demand)
    doc.add_heading('事实与理由：', level=1)
    doc.add_paragraph("此部分将由 AI 大模型基于提示词自动扩写...")

    doc.save(file_path)
    # 直接将生成的 Word 文件流返回给前端下载
    return send_file(file_path, as_attachment=True, download_name=f"{plaintiff}起诉状.docx")


if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5000, debug=True)